import os
import sys
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import ttkbootstrap as ttkb
from ttkbootstrap.constants import *
from ttkbootstrap.tooltip import ToolTip
import threading
import queue
import google.generativeai as genai
import speech_recognition as sr
from openai import OpenAI
import docx
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
import pygame
import pygame.mixer as mixer
import json
import socket
import subprocess
from cryptography.fernet import Fernet
import re
import warnings
from pathlib import Path
from pydub import AudioSegment
from pydub.effects import normalize
import math
import tempfile
import time

try:
    import whisper
except ImportError:
    whisper = None

warnings.filterwarnings("ignore", message="FP16 is not supported on CPU; using FP32 instead")

CONFIG_FILE = "voicevoyager_config.json"

def set_ffmpeg_path():
    try:
        subprocess.run(["ffmpeg", "-version"], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        print("FFmpeg found in system PATH.")
    except (subprocess.CalledProcessError, FileNotFoundError):
        base_path = os.path.dirname(os.path.abspath(__file__))
        if os.name == 'nt':
            ffmpeg_path = os.path.join(base_path, 'ffmpeg', 'win', 'ffmpeg.exe')
        else:
            ffmpeg_path = os.path.join(base_path, 'ffmpeg', 'mac' if os.uname().sysname == 'Darwin' else 'linux', 'ffmpeg')
            os.chmod(ffmpeg_path, 0o755)
        if not os.path.exists(ffmpeg_path):
            raise FileNotFoundError(f"FFmpeg not found at {ffmpeg_path}. Ensure it is bundled or installed.")
        AudioSegment.ffmpeg = ffmpeg_path
        print(f"Using bundled FFmpeg at {ffmpeg_path}")

set_ffmpeg_path()

os.environ["SDL_AUDIODRIVER"] = "directsound"
pygame.init()
mixer.init(buffer=32768, frequency=44100, channels=1)

class VoiceVoyager:
    def __init__(self, root):
        self.root = root
        self.root.title("skapezMpier VoiceVoyager")
        self.root.geometry("1000x860")
        self.style = ttkb.Style()
        self.style.theme_use("flatly")
        self.theme_var = tk.StringVar(value="flatly")

        # Dynamically resolve the path to icon.ico
        self.icon_path = self.resource_path("icon.ico")
        self.set_window_icon(self.root)

        self.center_window(self.root, 1000, 860)

        self.key = b'D2RU1VRyFjEQU24RSmfz8bQELCeaFfboNQUKxXKz6io='
        self.cipher = Fernet(self.key)

        self.api_config_file = "api_config.json"
        self.gemini_api_key = tk.StringVar(value="")
        self.openai_api_key = tk.StringVar(value="")
        self.load_api_keys()
        self.configure_apis()

        self.menu_bar = tk.Menu(self.root)
        self.root.config(menu=self.menu_bar)
        self.file_menu = tk.Menu(self.menu_bar, tearoff=0)
        self.menu_bar.add_cascade(label="File", menu=self.file_menu)
        self.file_menu.add_command(label="Open (Ctrl+O)", command=self.select_file)
        self.file_menu.add_command(label="Exit (Ctrl+Q)", command=self.root.quit)
        self.help_menu = tk.Menu(self.menu_bar, tearoff=0)
        self.menu_bar.add_cascade(label="Help", menu=self.help_menu)
        self.help_menu.add_command(label="Help", command=self.show_help)
        self.help_menu.add_command(label="About", command=self.show_about)
        self.license_menu = tk.Menu(self.menu_bar, tearoff=0)
        self.menu_bar.add_cascade(label="License", menu=self.license_menu)
        self.license_menu.add_command(label="View License", command=self.show_license)
        self.root.bind("<Control-o>", lambda e: self.select_file())
        self.root.bind("<Control-q>", lambda e: self.root.quit())

        self.main_frame = ttkb.Frame(root)
        self.main_frame.pack(fill=BOTH, expand=True)

        self.notebook = ttkb.Notebook(self.main_frame)
        self.notebook.pack(fill=BOTH, expand=True, padx=10, pady=10)

        self.transcript_tab = ttkb.Frame(self.notebook)
        self.analysis_tab = ttkb.Frame(self.notebook)
        self.api_tab = ttkb.Frame(self.notebook)
        self.notebook.add(self.transcript_tab, text="Transcription")
        self.notebook.add(self.analysis_tab, text="Analysis")
        self.notebook.add(self.api_tab, text="API")

        self.file_path = None
        self.audio = None
        self.audio_duration = 0
        self.full_transcription = ""
        self.transcription_thread = None
        self.is_playing = False
        self.is_transcribing = False
        self.is_processing = False
        self.sound_obj = None
        self.temp_file = None
        self.current_position = 0
        self.start_time_var = tk.DoubleVar(value=0)
        self.end_time_var = tk.DoubleVar(value=0)
        self.selected_duration_var = tk.StringVar(value="Selected Duration: 0s")
        self.current_position_var = tk.StringVar(value="Current: 0s")
        self.playback_queue = queue.Queue()

        self.setup_transcript_tab()
        self.setup_analysis_tab()
        self.setup_api_tab()

        self.update_text_colors()

        self.status_frame = ttkb.Frame(root)
        self.status_frame.pack(fill=X, side=BOTTOM)
        self.status_var = tk.StringVar(value="Ready")
        self.status_bar = ttkb.Label(self.status_frame, textvariable=self.status_var, bootstyle="inverse-success", padding=5)
        self.status_bar.pack(side=LEFT, fill=X, expand=True)
        self.spinner_var = tk.StringVar(value="")
        self.spinner_label = ttkb.Label(self.status_frame, textvariable=self.spinner_var, bootstyle="info")
        self.spinner_label.pack(side=RIGHT, padx=5)
        self.spinner_label.pack_forget()
        self.spinner_states = ["|", "/", "-", "\\"]
        self.spinner_index = 0

        self.load_preferences()
        if not self.hide_intro:
            self.show_intro_modal()

    def resource_path(self, relative_path):
        """Get the absolute path to a resource, works for dev and for PyInstaller"""
        try:
            # PyInstaller creates a temp folder and stores path in _MEIPASS
            base_path = sys._MEIPASS
        except AttributeError:
            base_path = os.path.abspath(".")
        return os.path.join(base_path, relative_path)

    def set_window_icon(self, window):
        """Set the icon for a given window"""
        try:
            window.iconbitmap(self.icon_path)
        except tk.TclError as e:
            print(f"Error setting icon: {e}")
    
    
    def show_custom_warning(self, title, message):
        dialog = tk.Toplevel(self.root)
        self.set_window_icon(dialog)  # Set the icon for the dialog
        dialog.title(title)
        self.center_window(dialog, 400, 150)
        dialog.transient(self.root)
        dialog.grab_set()

        ttkb.Label(dialog, text=message, wraplength=350).pack(pady=10)
        ttkb.Button(dialog, text="OK", bootstyle="primary", command=dialog.destroy).pack(pady=10)
        
    def on_dialog_close(self, dialog):
        self.is_processing = False
        self.update_playback_buttons()
        dialog.destroy()

    def center_window(self, window, width, height):
        screen_width = window.winfo_screenwidth()
        screen_height = window.winfo_screenheight()
        x = (screen_width - width) // 2
        y = (screen_height - height) // 2
        window.geometry(f"{width}x{height}+{x}+{y}")

    def load_api_keys(self):
        if os.path.exists(self.api_config_file):
            with open(self.api_config_file, "rb") as f:
                encrypted_data = f.read()
            try:
                decrypted_data = self.cipher.decrypt(encrypted_data).decode()
                config = json.loads(decrypted_data)
                self.gemini_api_key.set(config.get("gemini_api_key", ""))
                self.openai_api_key.set(config.get("openai_api_key", ""))
            except Exception:
                pass

    def save_api_keys(self):
        config = {"gemini_api_key": self.gemini_api_key.get(), "openai_api_key": self.openai_api_key.get()}
        encrypted_data = self.cipher.encrypt(json.dumps(config).encode())
        with open(self.api_config_file, "wb") as f:
            f.write(encrypted_data)
        self.configure_apis()
        messagebox.showinfo("API Keys", "API keys saved securely!")

    def configure_apis(self):
        gemini_key = self.gemini_api_key.get() or "YOUR_GEMINI_API_KEY"
        openai_key = self.openai_api_key.get() or "YOUR_OPENAI_API_KEY"
        genai.configure(api_key=gemini_key)
        self.gemini_model = genai.GenerativeModel("gemini-1.5-flash")
        self.openai_client = OpenAI(api_key=openai_key)

    def load_preferences(self):
        self.hide_intro = False
        if os.path.exists(CONFIG_FILE):
            with open(CONFIG_FILE, "r") as f:
                try:
                    config = json.load(f)
                    self.hide_intro = config.get("hide_intro", False)
                except json.JSONDecodeError:
                    pass

    def save_preferences(self, hide_intro):
        config = {"hide_intro": hide_intro}
        with open(CONFIG_FILE, "w") as f:
            json.dump(config, f)

    def setup_transcript_tab(self):
        self.file_frame = ttkb.LabelFrame(self.transcript_tab, text="Audio Input", padding=10)
        self.file_frame.pack(fill=X, pady=5)

        self.file_path_var = tk.StringVar(value="No file selected")
        self.file_path_label = ttkb.Label(self.file_frame, textvariable=self.file_path_var, wraplength=800)
        self.file_path_label.pack(side=LEFT, padx=5, fill=X, expand=True)

        self.browse_button = ttkb.Button(self.file_frame, text="Browse", bootstyle="success-outline", command=self.select_file)
        self.browse_button.pack(side=RIGHT, padx=5)
        ToolTip(self.browse_button, text="Select an audio file to transcribe")

        self.duration_label = ttkb.Label(self.file_frame, text="Duration: N/A")
        self.duration_label.pack(side=LEFT, padx=5)

        self.settings_frame = ttkb.LabelFrame(self.transcript_tab, text="Settings", padding=10)
        self.settings_frame.pack(fill=X, pady=5)

        self.model_var = tk.StringVar(value="Gemini")
        model_options = ["Gemini", "Google Speech", "Whisper", "PocketSphinx (Offline)"]
        if whisper:
            model_options.append("Whisper Local (Offline)")
        self.model_menu = ttkb.Combobox(self.settings_frame, textvariable=self.model_var, values=model_options, state="readonly", bootstyle="info")
        self.model_menu.pack(side=LEFT, padx=5)
        ToolTip(self.model_menu, text="Select the transcription model")

        self.language_var = tk.StringVar(value="en-US")
        self.languages = ["en-US", "es-ES", "fr-FR", "de-DE", "it-IT", "ja-JP", "zh-CN", "hi-IN", "ar-SA", "pt-BR", "ru-RU", "ko-KR"]
        self.language_menu = ttkb.Combobox(self.settings_frame, textvariable=self.language_var, values=self.languages, state="readonly", bootstyle="info")
        self.language_menu.pack(side=LEFT, fill=X, expand=True, padx=5)
        ToolTip(self.language_menu, text="Select the language of the audio")

        self.chunk_var = tk.IntVar(value=10)
        ttkb.Label(self.settings_frame, text="Chunk (s):").pack(side=LEFT, padx=5)
        self.chunk_spinbox = ttkb.Spinbox(self.settings_frame, from_=5, to=300, textvariable=self.chunk_var, width=5)
        self.chunk_spinbox.pack(side=LEFT, padx=5)
        ToolTip(self.chunk_spinbox, text="Set the duration of each transcription chunk (5-300 seconds)")

        self.normalize_var = tk.BooleanVar(value=True)
        self.normalize_check = ttkb.Checkbutton(self.settings_frame, text="Normalize Audio", variable=self.normalize_var, bootstyle="info")
        self.normalize_check.pack(side=LEFT, padx=5)
        ToolTip(self.normalize_check, text="Normalize audio volume for better transcription accuracy")

        self.control_frame = ttkb.Frame(self.transcript_tab)
        self.control_frame.pack(fill=X, pady=10)

        self.transcribe_button = ttkb.Button(self.control_frame, text="Transcribe", bootstyle="primary", command=self.start_transcription, state="disabled")
        self.transcribe_button.pack(side=LEFT, padx=5)
        ToolTip(self.transcribe_button, text="Start transcribing the selected audio")

        self.clear_button = ttkb.Button(self.control_frame, text="Clear", bootstyle="warning-outline", command=self.clear_output)
        self.clear_button.pack(side=LEFT, padx=5)
        ToolTip(self.clear_button, text="Clear the current transcription and reset the app")

        self.play_button = ttkb.Button(self.control_frame, text="Play", bootstyle="secondary-outline", command=self.play_audio, state="disabled")
        self.play_button.pack(side=LEFT, padx=5)
        ToolTip(self.play_button, text="Play the selected audio")

        self.stop_button = ttkb.Button(self.control_frame, text="Stop", bootstyle="danger-outline", command=self.stop_audio, state="disabled")
        self.stop_button.pack(side=LEFT, padx=5)
        ToolTip(self.stop_button, text="Stop audio playback and reset position")

        self.seek_scale = ttkb.Scale(self.control_frame, from_=0, to=0, orient=tk.HORIZONTAL, command=self.seek_audio)
        self.seek_scale.pack(side=LEFT, fill=X, expand=True, padx=5)
        ToolTip(self.seek_scale, text="Seek to a specific position in the audio")

        self.duration_display = ttkb.Label(self.control_frame, textvariable=self.selected_duration_var)
        self.duration_display.pack(side=LEFT, padx=5)

        self.current_position_label = ttkb.Label(self.control_frame, textvariable=self.current_position_var)
        self.current_position_label.pack(side=LEFT, padx=5)

        self.progress = ttkb.Progressbar(self.control_frame, mode="determinate", bootstyle="info", maximum=100)
        self.progress.pack(side=LEFT, fill=X, expand=True, padx=5)

        self.output_frame = ttkb.LabelFrame(self.transcript_tab, text="Transcription", padding=10)
        self.output_frame.pack(fill=BOTH, expand=True, pady=5)

        self.output_text = ttkb.ScrolledText(self.output_frame, height=15, wrap=tk.WORD, font=("Helvetica", 10))
        self.output_text.pack(fill=BOTH, expand=True)
        self.output_text.tag_config("error", foreground="red")

        ttkb.Label(self.transcript_tab, text="Theme:").pack(side=LEFT, padx=5)
        self.theme_combo = ttkb.Combobox(self.transcript_tab, textvariable=self.theme_var, values=["flatly", "darkly"], state="readonly", bootstyle="info")
        self.theme_combo.pack(side=LEFT, padx=5)
        self.theme_combo.bind("<<ComboboxSelected>>", self.switch_theme)
        ToolTip(self.theme_combo, text="Switch between light (flatly) and dark (darkly) themes")

    def setup_analysis_tab(self):
        self.analysis_frame = ttkb.LabelFrame(self.analysis_tab, text="Analysis Tools", padding=10)
        self.analysis_frame.pack(fill=X, pady=5)

        self.translate_button = ttkb.Button(self.analysis_frame, text="Translate", bootstyle="info-outline", command=self.translate_transcript, state="disabled")
        self.translate_button.pack(side=LEFT, padx=5)
        ToolTip(self.translate_button, text="Translate the transcription to another language")

        self.keyword_button = ttkb.Button(self.analysis_frame, text="Extract Keywords", bootstyle="info-outline", command=self.extract_keywords)
        self.keyword_button.pack(side=LEFT, padx=5)
        ToolTip(self.keyword_button, text="Extract key phrases or words from the transcription")

        self.action_button = ttkb.Button(self.analysis_frame, text="Detect Actions", bootstyle="info-outline", command=self.detect_action_items)
        self.action_button.pack(side=LEFT, padx=5)
        ToolTip(self.action_button, text="Identify action items or tasks in the transcription")

        self.event_button = ttkb.Button(self.analysis_frame, text="Tag Events", bootstyle="info-outline", command=self.tag_audio_events)
        self.event_button.pack(side=LEFT, padx=5)
        ToolTip(self.event_button, text="Tag non-speech events (e.g., laughter, applause)")

        self.qa_button = ttkb.Button(self.analysis_frame, text="Ask Question", bootstyle="info-outline", command=self.ask_question)
        self.qa_button.pack(side=LEFT, padx=5)
        ToolTip(self.qa_button, text="Ask a question about the transcription")

        self.sentiment_button = ttkb.Button(self.analysis_frame, text="Sentiment Analysis", bootstyle="info-outline", command=self.sentiment_analysis)
        self.sentiment_button.pack(side=LEFT, padx=5)
        ToolTip(self.sentiment_button, text="Analyze the sentiment and tone of the transcription")

        self.export_button = ttkb.Button(self.analysis_frame, text="Export", bootstyle="info-outline", command=self.export_transcription)
        self.export_button.pack(side=LEFT, padx=5)
        ToolTip(self.export_button, text="Export the transcription as PDF or DOCX")

        self.analysis_output = ttkb.ScrolledText(self.analysis_tab, height=15, wrap=tk.WORD, font=("Helvetica", 10))
        self.analysis_output.pack(fill=BOTH, expand=True, pady=5)
        self.analysis_output.tag_config("keyword", foreground="darkblue")
        self.analysis_output.tag_config("action", foreground="darkgreen")
        self.analysis_output.tag_config("sentiment", foreground="darkmagenta")

    def setup_api_tab(self):
        self.api_frame = ttkb.LabelFrame(self.api_tab, text="API Keys", padding=10)
        self.api_frame.pack(fill=BOTH, expand=True, pady=5)

        ttkb.Label(self.api_frame, text="Gemini API Key:").grid(row=0, column=0, padx=5, pady=5, sticky="w")
        ttkb.Entry(self.api_frame, textvariable=self.gemini_api_key, width=50).grid(row=0, column=1, padx=5, pady=5)
        ttkb.Label(self.api_frame, text="OpenAI API Key:").grid(row=1, column=0, padx=5, pady=5, sticky="w")
        ttkb.Entry(self.api_frame, textvariable=self.openai_api_key, width=50).grid(row=1, column=1, padx=5, pady=5)
        self.save_api_button = ttkb.Button(self.api_frame, text="Save API Keys", bootstyle="success-outline", command=self.save_api_keys)
        self.save_api_button.grid(row=2, column=1, padx=5, pady=10, sticky="e")
        ToolTip(self.save_api_button, text="Save the API keys securely")

    def switch_theme(self, event):
        self.style.theme_use(self.theme_var.get())
        self.status_bar.configure(bootstyle="inverse-success" if self.theme_var.get() == "flatly" else "inverse-dark")
        self.update_text_colors()

    def update_text_colors(self):
        default_color = "white" if self.theme_var.get() == "darkly" else "black"
        self.output_text.configure(foreground=default_color)
        self.analysis_output.configure(foreground=default_color)
        if self.theme_var.get() == "flatly":
            self.analysis_output.tag_config("keyword", foreground="darkblue")
            self.analysis_output.tag_config("action", foreground="darkgreen")
            self.analysis_output.tag_config("sentiment", foreground="darkmagenta")
        else:
            self.analysis_output.tag_config("keyword", foreground="lightblue")
            self.analysis_output.tag_config("action", foreground="lightgreen")
            self.analysis_output.tag_config("sentiment", foreground="violet")

    def show_help(self):
        # Removed unnecessary line as 'dialog' is not defined
        help_text = (
            "VoiceVoyager Help\n\n"
            "1. Transcription Tab:\n"
            "   - Browse: Select an audio file (.mp3, .wav, .aiff, .flac).\n"
            "   - Model: Choose Gemini, Google Speech, or Whisper.\n"
            "   - Language: Select the audio language.\n"
            "   - Chunk: Set how long each transcription segment is (5-300s).\n"
            "   - Normalize: Adjust volume for better accuracy.\n"
            "   - Play/Stop: Control audio playback.\n"
            "   - Seek: Move to a specific part of the audio.\n"
            "   - Transcribe: Start transcription.\n\n"
            "2. Analysis Tab:\n"
            "   - Translate: Convert transcription to another language.\n"
            "   - Extract Keywords: Find important words/phrases.\n"
            "   - Detect Actions: List tasks or action items.\n"
            "   - Tag Events: Mark non-speech sounds (e.g., laughter).\n"
            "   - Ask Question: Query the transcription.\n"
            "   - Sentiment Analysis: Analyze tone with emojis.\n"
            "   - Export: Save as PDF or DOCX.\n\n"
            "3. API Tab:\n"
            "   - Enter and save encrypted API keys for Gemini and OpenAI.\n\n"
            "Shortcuts:\n"
            "   - Ctrl+O: Open file\n"
            "   - Ctrl+Q: Exit"
        )
        dialog = tk.Toplevel(self.root)
        self.set_window_icon(dialog)  # Set the icon for the dialog
        dialog.title("Help")
        self.center_window(dialog, 800, 600)
        dialog.transient(self.root)
        dialog.grab_set()

        content_frame = ttkb.Frame(dialog)
        content_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        text_frame = ttkb.Frame(content_frame)
        text_frame.pack(fill=tk.BOTH, expand=True)

        help_text_widget = ttkb.ScrolledText(text_frame, wrap=tk.WORD, height=15, font=("Helvetica", 10))
        help_text_widget.insert(tk.END, help_text)
        help_text_widget.configure(state="disabled")
        help_text_widget.pack(fill=tk.BOTH, expand=True)

        ttkb.Button(content_frame, text="OK", bootstyle="primary", command=dialog.destroy).pack(pady=5)
        # messagebox.showinfo("Help", help_text)

    def show_about(self):
        about_text = (
            "VoiceVoyager v1.0\n\n"
            "Created by skapezMpier\n\n"
            "About VoiceVoyager:\n"
            "VoiceVoyager is a comprehensive audio transcription and analysis tool designed to simplify converting audio into actionable insights.\n\n"
            "Features:\n"
            "- Supports Gemini, Google Speech, Whisper, PocketSphinx, and Whisper Local.\n"
            "- Tools for keywords, actions, events, sentiment, and export.\n"
            "- Interactive playback with seek.\n"
            "- Secure API key management.\n\n"
            "Contact: skapezmpier@gmail.com"
        )
        dialog = tk.Toplevel(self.root)
        self.set_window_icon(dialog)  # Set the icon for the dialog
        dialog.title("About")
        self.center_window(dialog, 800, 500)
        dialog.transient(self.root)
        dialog.grab_set()

        content_frame = ttkb.Frame(dialog)
        content_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        text_frame = ttkb.Frame(content_frame)
        text_frame.pack(fill=tk.BOTH, expand=True)

        about_text_widget = ttkb.ScrolledText(text_frame, wrap=tk.WORD, height=10, font=("Helvetica", 10))
        about_text_widget.insert(tk.END, about_text)
        about_text_widget.configure(state="disabled")
        about_text_widget.pack(fill=tk.BOTH, expand=True)

        ttkb.Button(content_frame, text="OK", bootstyle="primary", command=dialog.destroy).pack(pady=5)
        # messagebox.showinfo("About", about_text)

    def show_license(self):
        license_text = (
            "VoiceVoyage Technology Software License\n\n"
            "Non-Commercial Open Source License\n\n"
            "Version 1.0\n"
            "Copyright (c) 2025 skapezMpier\n\n"
            "1. Permission Grant\n"
            "This software is released as open source. You are permitted to:\n"
            "- Use, modify, and distribute the software for personal, research, or integration into commercial products, provided that the software itself remains freely accessible.\n"
            "- Contribute improvements and modifications, provided they are shared under the same license.\n"
            "- Share this software with others under these same terms.\n\n"
            "2. Accessibility Requirement\n"
            "- This software may be included in commercial products, but it must remain free to access and use within those products.\n"
            "- You may not restrict access to this software behind a paywall, subscription, or exclusive service model.\n\n"
            "3. Attribution\n"
            "You must provide credit to the original author(s) by maintaining this license notice in all copies and derivatives of the software.\n\n"
            "4. Disclaimer\n"
            "THIS SOFTWARE IS PROVIDED \"AS IS\", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR IMPLIED. THE AUTHOR(S) SHALL NOT BE LIABLE FOR ANY CLAIMS OR DAMAGES ARISING FROM ITS USE.\n\n"
            "5. License Compatibility\n"
            "If you contribute to this project, you agree that your contributions will be licensed under the same terms. Any modifications must also follow this accessibility requirement.\n\n"
            "For further inquiries regarding licensing, contact:\n"
            "    - skapezempire@gmail.com\n"
            "   - https://www.linkedin.com/company/skapezmpier/"
        )
        dialog = tk.Toplevel(self.root)
        self.set_window_icon(dialog)  # Set the icon for the dialog
        dialog.title("License")
        self.center_window(dialog, 800, 600)
        dialog.transient(self.root)
        dialog.grab_set()

        content_frame = ttkb.Frame(dialog)
        content_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        text_frame = ttkb.Frame(content_frame)
        text_frame.pack(fill=tk.BOTH, expand=True)

        license_text_widget = ttkb.ScrolledText(text_frame, wrap=tk.WORD, height=15, font=("Helvetica", 10))
        license_text_widget.insert(tk.END, license_text)
        license_text_widget.configure(state="disabled")
        license_text_widget.pack(fill=tk.BOTH, expand=True)

        ttkb.Button(content_frame, text="OK", bootstyle="primary", command=dialog.destroy).pack(pady=5)
            
        # messagebox.showinfo("License", license_text)

    def show_intro_modal(self):
        intro_text = (
            "🎙️ Welcome to VoiceVoyager! 🎙️\n\n"
            "VoiceVoyager is your ultimate audio transcription and analysis tool, designed to make your audio processing seamless and efficient.\n\n"
            "✨ Key Features:\n"
            "- Transcribe audio using online models (Gemini, Google Speech, Whisper) or offline models (PocketSphinx, Whisper Local).\n"
            "- Analyze transcriptions for keywords, action items, sentiment, and more.\n"
            "- Interactive audio playback with play and stop controls.\n"
            "- Export transcriptions as PDF or DOCX files.\n\n"
            "📋 Requirements:\n"
            "- Internet connection for online transcription models (Gemini, Google Speech, Whisper).\n"
            "- FFmpeg installed (via system PATH or bundled with the app).\n"
            "- API keys for Gemini and OpenAI (optional, for online transcription and analysis).\n"
            "- Offline Whisper model (optional, auto-downloads if selected).\n\n"
            "🚀 Getting Started:\n"
            "- Select an audio file in the Transcription tab.\n"
            "- Choose your preferred transcription model and language.\n"
            "- Click 'Transcribe' to start processing.\n"
            "- Use the Analysis tab for deeper insights into your transcription.\n"
            "- Configure API keys in the API tab if using online models.\n\n"
            "Click OK to dive in!"
        )
        dialog = tk.Toplevel(self.root)
        self.set_window_icon(dialog)
        dialog.title("Welcome to VoiceVoyager")
        dialog_width, dialog_height = 600, 600
        self.center_window(dialog, dialog_width, dialog_height)

        dialog.transient(self.root)
        dialog.grab_set()

        content_frame = ttkb.Frame(dialog)
        content_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        text_frame = ttkb.Frame(content_frame)
        text_frame.pack(fill=tk.BOTH, expand=True)

        intro_text_widget = ttkb.ScrolledText(text_frame, wrap=tk.WORD, height=15, font=("Helvetica", 10))
        intro_text_widget.insert(tk.END, intro_text)
        intro_text_widget.configure(state="disabled")
        intro_text_widget.pack(fill=tk.BOTH, expand=True)

        bottom_frame = ttkb.Frame(content_frame)
        bottom_frame.pack(fill=X, pady=5)

        hide_var = tk.BooleanVar(value=False)
        ttkb.Checkbutton(bottom_frame, text="Don't show this again", variable=hide_var, bootstyle="info").pack(side=tk.LEFT, pady=5)

        ttkb.Button(bottom_frame, text="OK", bootstyle="primary", command=lambda: self.on_intro_ok(hide_var, dialog)).pack(side=tk.RIGHT, pady=5)
        ToolTip(bottom_frame, text="Click to close the introduction")

    def on_intro_ok(self, hide_var, dialog):
        if hide_var.get():
            self.save_preferences(True)
        dialog.destroy()

    def check_internet(self):
        try:
            socket.create_connection(("www.google.com", 80), timeout=2)
            return True
        except OSError:
            return False

    def select_file(self):
        self.stop_audio()
        self.file_path = filedialog.askopenfilename(filetypes=[("Audio Files", "*.mp3 *.wav *.aiff *.flac"), ("All Files", "*.*")])
        if self.file_path:
            self.file_path_var.set(self.file_path)
            self.transcribe_button.config(state="normal")
            self.play_button.config(state="normal" if not self.is_transcribing else "disabled")
            self.output_text.delete(1.0, tk.END)
            self.output_text.insert(tk.END, f"Selected file: {self.file_path}\n")
            self.status_var.set("File selected")
            self.audio = AudioSegment.from_file(self.file_path)
            self.audio_duration = len(self.audio) / 1000
            self.duration_label.config(text=f"Duration: {self.audio_duration:.1f} seconds")
            self.seek_scale.config(to=self.audio_duration)
            self.end_time_var.set(self.audio_duration)
            default_chunk = max(5, min(300, int(self.audio_duration / 10)))
            self.chunk_var.set(default_chunk)
            self.update_selected_duration()
            self.full_transcription = ""
            if self.model_var.get() == "Gemini" and self.audio_duration > 900:
                self.show_custom_warning("Gemini Limit", "Audio exceeds 15 minutes. Adjust chunk size or switch models.")
            self.update_playback_buttons()

    def play_audio(self):
        if self.is_transcribing or self.is_processing:
            self.show_custom_warning("Processing in Progress", "Cannot play audio while transcription or analysis is processing.")
            return
        if self.audio and not self.is_playing:
            self.stop_audio()  # Ensure previous playback is stopped
            try:
                # Standardize audio format for compatibility
                sample_rate = 44100
                channels = 1  # Mono for simplicity
                self.audio = self.audio.set_frame_rate(sample_rate).set_channels(channels)

                # Create and export temp file
                self.temp_file = tempfile.NamedTemporaryFile(suffix=".wav", delete=False).name
                self.audio.export(self.temp_file, format="wav", parameters=["-ar", str(sample_rate), "-ac", str(channels)])

                # Verify temp file
                if not os.path.exists(self.temp_file):
                    raise Exception(f"Temp file {self.temp_file} not created")

                # Use Sound object for playback
                self.sound_obj = mixer.Sound(self.temp_file)
                self.sound_obj.play()
                self.is_playing = True
                self.update_playback_buttons()
            except Exception as e:
                self.output_text.insert(tk.END, f"Error playing audio: {e}\n", "error")
                self.status_var.set(f"Error: {e}")

    def stop_audio(self):
        self.is_playing = False
        self.current_position = 0
        if self.sound_obj:
            self.sound_obj.stop()
            self.sound_obj = None
        if self.temp_file and os.path.exists(self.temp_file):
            try:
                os.unlink(self.temp_file)
            except Exception as e:
                print(f"Error deleting temp file: {e}")
            self.temp_file = None
        self.seek_scale.set(0)
        self.current_position_var.set("Current: 0s")
        self.status_var.set("Audio stopped")
        self.update_playback_buttons()

    def seek_audio(self, value):
        self.current_position = float(value)
        self.current_position_var.set(f"Current: {self.current_position:.1f}s")
        self.start_time_var.set(self.current_position)
        self.update_selected_duration()
        if self.is_playing and self.sound_obj:
            self.sound_obj.stop()
            self.sound_obj.play(start=self.current_position)
            self.root.after(100, self.monitor_playback)

    def monitor_playback(self):
        if self.is_playing and self.sound_obj:
            pos = self.current_position + (pygame.mixer.get_pos() / 1000.0)
            if pos >= self.audio_duration:
                self.stop_audio()
            else:
                self.seek_scale.set(pos)
                self.current_position_var.set(f"Current: {pos:.1f}s")
                self.root.after(100, self.monitor_playback)

    def update_playback_buttons(self):
        disable_all = self.is_transcribing or self.is_processing
        self.play_button.config(state="normal" if self.audio and not self.is_playing and not disable_all else "disabled")
        self.stop_button.config(state="normal" if self.is_playing and not disable_all else "disabled")
        self.transcribe_button.config(state="normal" if self.file_path and not disable_all else "disabled")
        self.clear_button.config(state="normal" if not disable_all else "disabled")
        self.translate_button.config(state="normal" if self.full_transcription and not disable_all else "disabled")
        self.keyword_button.config(state="normal" if self.full_transcription and not disable_all else "disabled")
        self.action_button.config(state="normal" if self.full_transcription and not disable_all else "disabled")
        self.event_button.config(state="normal" if self.full_transcription and not disable_all else "disabled")
        self.qa_button.config(state="normal" if self.full_transcription and not disable_all else "disabled")
        self.sentiment_button.config(state="normal" if self.full_transcription and not disable_all else "disabled")
        self.export_button.config(state="normal" if self.full_transcription and not disable_all else "disabled")
        self.browse_button.config(state="normal" if not disable_all else "disabled")
        self.model_menu.config(state="readonly" if not disable_all else "disabled")
        self.language_menu.config(state="readonly" if not disable_all else "disabled")
        self.chunk_spinbox.config(state="normal" if not disable_all else "disabled")
        self.normalize_check.config(state="normal" if not disable_all else "disabled")
        self.seek_scale.config(state="normal" if self.audio and not disable_all else "disabled")
        self.theme_combo.config(state="readonly" if not disable_all else "disabled")
        self.save_api_button.config(state="normal" if not disable_all else "disabled")

    def update_selected_duration(self):
        duration = self.end_time_var.get() - self.start_time_var.get()
        self.selected_duration_var.set(f"Selected Duration: {duration:.1f}s")

    def animate_spinner(self):
        if self.is_transcribing or self.is_processing:
            self.spinner_index = (self.spinner_index + 1) % len(self.spinner_states)
            self.spinner_var.set(self.spinner_states[self.spinner_index])
            self.root.after(100, self.animate_spinner)

    def start_transcription(self):
        if self.model_var.get() not in ["PocketSphinx (Offline)", "Whisper Local (Offline)"] and not self.check_internet():
            messagebox.showerror("No Internet", "Internet required for online transcription.")
            return
        if self.model_var.get() == "Whisper Local (Offline)" and not whisper:
            messagebox.showerror("Whisper Missing", "Install 'openai-whisper' for offline Whisper: pip install openai-whisper")
            return
        self.is_transcribing = True
        self.update_playback_buttons()
        self.output_text.delete(1.0, tk.END)
        self.output_text.insert(tk.END, "Starting transcription...\n")
        self.status_var.set("Processing")
        self.spinner_label.pack(side=RIGHT, padx=5)
        self.animate_spinner()
        self.progress["value"] = 0
        self.transcription_thread = threading.Thread(target=self.transcribe_audio, daemon=True)
        self.transcription_thread.start()

    def clean_text(self, text):
        text = re.sub(r'\*{1,3}(.*?)\*{1,3}', r'\1', text)
        text = re.sub(r'_{1,2}(.*?)_{1,2}', r'\1', text)
        text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
        return text.strip()

    def transcribe_audio(self):
        if not os.path.exists(self.file_path):
            self.output_text.insert(tk.END, f"Error: File '{self.file_path}' does not exist.\n", "error")
            self.finish_transcription(error=True)
            return

        audio = self.audio
        if self.normalize_var.get():
            audio = normalize(audio.set_channels(1).set_frame_rate(16000))
        total_duration = len(audio) / 1000
        self.output_text.insert(tk.END, f"Total audio duration: {total_duration} seconds\n")

        start_time = self.start_time_var.get() * 1000
        end_time = self.end_time_var.get() * 1000 if self.end_time_var.get() > 0 else total_duration * 1000
        audio = audio[start_time:end_time]
        temp_file = tempfile.mktemp(suffix=".wav")
        audio.export(temp_file, format="wav")

        chunk_duration = self.chunk_var.get()
        num_chunks = math.ceil((end_time - start_time) / 1000 / chunk_duration)
        self.full_transcription = []
        start_time_process = time.time()

        model = self.model_var.get()
        for i in range(num_chunks):
            chunk_start = start_time + i * chunk_duration * 1000
            chunk_end = min(start_time + (i + 1) * chunk_duration * 1000, end_time)
            chunk = audio[chunk_start - start_time:chunk_end - start_time]
            chunk_file = tempfile.mktemp(suffix=".wav")
            chunk.export(chunk_file, format="wav")

            try:
                if model == "Gemini":
                    chunk_text = self.clean_text(self.transcribe_with_gemini(chunk_file, chunk_start, chunk_end))
                elif model == "Google Speech":
                    chunk_text = self.clean_text(self.transcribe_with_google(chunk_file, chunk_start, chunk_end))
                elif model == "Whisper":
                    chunk_text = self.clean_text(self.transcribe_with_whisper(chunk_file, chunk_start, chunk_end))
                elif model == "PocketSphinx (Offline)":
                    chunk_text = self.clean_text(self.transcribe_with_pocketsphinx(chunk_file, chunk_start, chunk_end))
                else:  # Whisper Local (Offline)
                    chunk_text = self.clean_text(self.transcribe_with_whisper_local(chunk_file, chunk_start, chunk_end))
                self.full_transcription.append(f"[{chunk_start/1000:.1f}-{chunk_end/1000:.1f}] {chunk_text}")
                self.output_text.insert(tk.END, f"[{chunk_start/1000:.1f}-{chunk_end/1000:.1f}] {chunk_text}\n")
                elapsed = time.time() - start_time_process
                remaining = (elapsed / (i + 1)) * (num_chunks - (i + 1))
                self.status_var.set(f"Processing ({i+1}/{num_chunks}, ETA: {remaining:.1f}s)")
                self.progress["value"] = ((i + 1) / num_chunks) * 100
                self.root.update_idletasks()
            except Exception as e:
                self.output_text.insert(tk.END, f"[{chunk_start/1000:.1f}-{chunk_end/1000:.1f}] [Error: {e}]\n", "error")
                self.finish_transcription(error=True)
                break
            finally:
                os.remove(chunk_file)
        os.remove(temp_file)
        self.finish_transcription()

    def transcribe_with_gemini(self, chunk_file, start, end):
        with open(chunk_file, "rb") as f:
            audio_data = f.read()
        prompt = f"Transcribe this audio in {self.language_var.get()} with speaker labels."
        response = self.gemini_model.generate_content([prompt, {"mime_type": "audio/wav", "data": audio_data}])
        return response.text

    def transcribe_with_google(self, chunk_file, start, end):
        r = sr.Recognizer()
        try:
            with sr.AudioFile(chunk_file) as source:
                audio = r.record(source)
            return r.recognize_google(audio, language=self.language_var.get())
        except sr.UnknownValueError:
            return "[Error: Google Speech could not understand the audio]"
        except sr.RequestError as e:
            return f"[Error: Google Speech API request failed - {e}]"

    def transcribe_with_whisper(self, chunk_file, start, end):
        with open(chunk_file, "rb") as f:
            response = self.openai_client.audio.transcriptions.create(model="whisper-1", file=f, language=self.language_var.get().split("-")[0])
        return response.text

    def transcribe_with_pocketsphinx(self, chunk_file, start, end):
        r = sr.Recognizer()
        with sr.AudioFile(chunk_file) as source:
            audio = r.record(source)
        return r.recognize_sphinx(audio, language="en-US")

    def transcribe_with_whisper_local(self, chunk_file, start, end):
        model = whisper.load_model("base")
        result = model.transcribe(chunk_file, language=self.language_var.get().split("-")[0])
        return result["text"]

    def finish_transcription(self, error=False):
        self.is_transcribing = False
        self.spinner_label.pack_forget()
        if not error:
            self.full_transcription = "\n".join(self.full_transcription)
            output_file = os.path.splitext(self.file_path)[0] + f"_{self.model_var.get().lower().replace(' ', '_')}.txt"
            with open(output_file, "w", encoding="utf-8") as f:
                f.write(self.full_transcription)
            self.output_text.insert(tk.END, f"\nTranscription saved to '{output_file}'\n")
            self.status_var.set("Completed")
        else:
            self.status_var.set("Error occurred")
        self.update_playback_buttons()

    def translate_transcript(self):
        if not self.full_transcription:
            self.show_custom_warning("No Transcription", "Please transcribe an audio file first.")
            return
        self.is_processing = True
        self.update_playback_buttons()
        self.status_var.set("Translating...")
        self.spinner_label.pack(side=RIGHT, padx=5)
        self.animate_spinner()

        dialog = tk.Toplevel(self.root)
        self.set_window_icon(dialog) 
        dialog.title("Translate")
        self.center_window(dialog, 400, 200)
        dialog.transient(self.root)
        dialog.grab_set()

        ttkb.Label(dialog, text="Select target language:").pack(pady=10)
        target_lang_var = tk.StringVar(value="es-ES")
        ttkb.Combobox(dialog, textvariable=target_lang_var, values=self.languages, state="readonly").pack(pady=10)
        ttkb.Button(dialog, text="Translate", bootstyle="primary", command=lambda: self.perform_translation(target_lang_var.get(), dialog)).pack(pady=10)

    def perform_translation(self, target_lang, dialog):
        
        def do_translation():
            try:
                response = self.gemini_model.generate_content(f"Translate this text to {target_lang}:\n\n{self.full_transcription}")
                cleaned_text = self.clean_text(response.text)
                self.output_text.delete(1.0, tk.END)
                self.output_text.insert(tk.END, f"Translated to {target_lang}:\n{cleaned_text}\n")
                self.status_var.set(f"Translated to {target_lang}")
            except Exception as e:
                self.output_text.insert(tk.END, f"Error during translation: {e}\n", "error")
                self.status_var.set("Translation Failed")
            finally:
                self.is_processing = False
                self.spinner_label.pack_forget()
                self.update_playback_buttons()
                dialog.destroy()
        threading.Thread(target=do_translation, daemon=True).start()

    def extract_keywords(self):
        if not self.full_transcription:
            self.show_custom_warning("No Transcription", "Please transcribe an audio file first.")
            return
        self.is_processing = True
        self.update_playback_buttons()
        self.status_var.set("Extracting Keywords...")
        self.spinner_label.pack(side=RIGHT, padx=5)
        self.animate_spinner()

        def do_extraction():
            try:
                response = self.gemini_model.generate_content(f"Extract key phrases or words from this text:\n\n{self.full_transcription}")
                cleaned_text = self.clean_text(response.text)
                self.analysis_output.delete(1.0, tk.END)
                self.analysis_output.insert(tk.END, "Key Phrases/Words:\n" + cleaned_text + "\n", "keyword")
                self.status_var.set("Keywords Extracted")
            except Exception as e:
                self.analysis_output.insert(tk.END, f"Error: {e}\n", "error")
                self.status_var.set("Keyword Extraction Failed")
            finally:
                self.is_processing = False
                self.spinner_label.pack_forget()
                self.update_playback_buttons()
        threading.Thread(target=do_extraction, daemon=True).start()

    def detect_action_items(self):
        if not self.full_transcription:
            
            self.show_custom_warning("No Transcription", "Please transcribe an audio file first.")
            return
        self.is_processing = True
        self.update_playback_buttons()
        self.status_var.set("Detecting Actions...")
        self.spinner_label.pack(side=RIGHT, padx=5)
        self.animate_spinner()

        def do_detection():
            try:
                response = self.gemini_model.generate_content(f"Identify action items or tasks from this text:\n\n{self.full_transcription}")
                cleaned_text = self.clean_text(response.text)
                self.analysis_output.delete(1.0, tk.END)
                self.analysis_output.insert(tk.END, "Action Items:\n" + cleaned_text + "\n", "action")
                self.status_var.set("Actions Detected")
            except Exception as e:
                self.analysis_output.insert(tk.END, f"Error: {e}\n", "error")
                self.status_var.set("Action Detection Failed")
            finally:
                self.is_processing = False
                self.spinner_label.pack_forget()
                self.update_playback_buttons()
        threading.Thread(target=do_detection, daemon=True).start()

    def tag_audio_events(self):
        
        if not self.full_transcription:
            
            self.show_custom_warning("No Transcription", "Please transcribe an audio file first.")
            return
        self.is_processing = True
        self.update_playback_buttons()
        self.status_var.set("Tagging Events...")
        self.spinner_label.pack(side=RIGHT, padx=5)
        self.animate_spinner()

        def do_tagging():
            try:
                response = self.gemini_model.generate_content(f"Tag non-speech events (e.g., laughter, applause) in this transcription:\n\n{self.full_transcription}")
                cleaned_text = self.clean_text(response.text)
                self.analysis_output.delete(1.0, tk.END)
                self.analysis_output.insert(tk.END, "Transcription with Events:\n" + cleaned_text + "\n")
                self.status_var.set("Events Tagged")
            except Exception as e:
                self.analysis_output.insert(tk.END, f"Error: {e}\n", "error")
                self.status_var.set("Event Tagging Failed")
            finally:
                self.is_processing = False
                self.spinner_label.pack_forget()
                self.update_playback_buttons()
        threading.Thread(target=do_tagging, daemon=True).start()

    def ask_question(self):
        if not self.full_transcription:
            self.show_custom_warning("No Transcription", "Please transcribe an audio file first.")
            return
        self.is_processing = True
        self.update_playback_buttons()

        dialog = tk.Toplevel(self.root)
        self.set_window_icon(dialog)  # Set the icon for the dialog
        dialog.title("Ask Question")
        self.center_window(dialog, 400, 200)
        dialog.transient(self.root)
        dialog.grab_set()

        # Reset is_processing when the dialog is closed
        dialog.protocol("WM_DELETE_WINDOW", lambda: self.on_dialog_close(dialog))

        ttkb.Label(dialog, text="Enter your question:").pack(pady=10)
        question_var = tk.StringVar()
        ttkb.Entry(dialog, textvariable=question_var).pack(pady=10)
        ttkb.Button(dialog, text="Ask", bootstyle="primary", command=lambda: self.perform_ask_question(question_var.get(), dialog)).pack(pady=10)

    def perform_ask_question(self, question, dialog):
        if question:
            self.status_var.set("Processing Question...")
            self.spinner_label.pack(side=RIGHT, padx=5)
            self.animate_spinner()

            def do_question():
                try:
                    response = self.gemini_model.generate_content(f"Answer this question based on the text:\n\n{self.full_transcription}\n\nQuestion: {question}")
                    cleaned_text = self.clean_text(response.text)
                    self.analysis_output.delete(1.0, tk.END)
                    self.analysis_output.insert(tk.END, f"Q: {question}\nA: {cleaned_text}\n")
                    self.status_var.set("Question Answered")
                except Exception as e:
                    self.analysis_output.insert(tk.END, f"Error: {e}\n", "error")
                    self.status_var.set("Question Processing Failed")
                finally:
                    self.is_processing = False
                    self.spinner_label.pack_forget()
                    self.update_playback_buttons()
                    dialog.destroy()
            threading.Thread(target=do_question, daemon=True).start()
        else:
            self.is_processing = False
            self.update_playback_buttons()
            dialog.destroy()

    def sentiment_analysis(self):
        if not self.full_transcription:
            self.show_custom_warning("No Transcription", "Please transcribe an audio file first.")
            return
        self.is_processing = True
        self.update_playback_buttons()
        self.status_var.set("Analyzing Sentiment...")
        self.spinner_label.pack(side=RIGHT, padx=5)
        self.animate_spinner()

        def do_sentiment():
            try:
                response = self.gemini_model.generate_content(f"Analyze the sentiment and tone of this text (e.g., confident, sad, happy):\n\n{self.full_transcription}")
                cleaned_text = self.clean_text(response.text)
                sentiment_with_emojis = self.add_emojis(cleaned_text)
                self.analysis_output.delete(1.0, tk.END)
                self.analysis_output.insert(tk.END, "Sentiment Analysis:\n" + sentiment_with_emojis + "\n", "sentiment")
                self.status_var.set("Sentiment Analyzed")
            except Exception as e:
                self.analysis_output.insert(tk.END, f"Error: {e}\n", "error")
                self.status_var.set("Sentiment Analysis Failed")
            finally:
                self.is_processing = False
                self.spinner_label.pack_forget()
                self.update_playback_buttons()
        threading.Thread(target=do_sentiment, daemon=True).start()

    def add_emojis(self, text):
        emoji_map = {
            "happy": "😊", "sad": "😢", "confident": "💪", "angry": "😠", "neutral": "😐",
            "excited": "🎉", "nervous": "😬", "calm": "😌", "surprised": "😲"
        }
        for tone, emoji in emoji_map.items():
            text = re.sub(rf'\b{tone}\b', f"{tone} {emoji}", text, flags=re.IGNORECASE)
        return text

    def export_transcription(self):
        if not self.full_transcription:
            self.show_custom_warning("No Transcription", "Please transcribe an audio file first.")
            return
        self.is_processing = True
        self.update_playback_buttons()

        dialog = tk.Toplevel(self.root)
        self.set_window_icon(dialog)  # Set the icon for the dialog
        dialog.title("Export")
        self.center_window(dialog, 400, 200)
        dialog.transient(self.root)
        dialog.grab_set()

        # Reset is_processing when the dialog is closed
        dialog.protocol("WM_DELETE_WINDOW", lambda: self.on_dialog_close(dialog))

        ttkb.Label(dialog, text="Enter file type (pdf/docx):").pack(pady=10)
        file_type_var = tk.StringVar()
        ttkb.Entry(dialog, textvariable=file_type_var).pack(pady=10)
        ttkb.Button(dialog, text="Export", bootstyle="primary", command=lambda: self.perform_export(file_type_var.get(), dialog)).pack(pady=10)

    def perform_export(self, file_type, dialog):
        # Validate file_type before starting the thread
        file_type = file_type.lower().strip()
        if file_type not in ["pdf", "docx"]:
            error_dialog = tk.Toplevel(self.root)
            self.set_window_icon(error_dialog)
            error_dialog.title("Invalid Type")
            self.center_window(error_dialog, 400, 150)
            error_dialog.transient(self.root)
            error_dialog.grab_set()
            ttkb.Label(error_dialog, text="Please enter 'pdf' or 'docx'.").pack(pady=10)
            ttkb.Button(error_dialog, text="OK", bootstyle="primary", command=error_dialog.destroy).pack(pady=10)
            self.is_processing = False
            self.update_playback_buttons()
            dialog.destroy()
            return

        self.status_var.set("Exporting...")
        self.spinner_label.pack(side=RIGHT, padx=5)
        self.animate_spinner()

        def do_export():
            try:
                output_file = os.path.splitext(self.file_path)[0] + f"_transcription.{file_type}"
                if file_type == "pdf":
                    doc = SimpleDocTemplate(output_file, pagesize=letter)
                    styles = getSampleStyleSheet()
                    story = [
                        Paragraph(f"Transcription of {os.path.basename(self.file_path)}", styles["Title"]),
                        Spacer(1, 12),
                        Paragraph(f"Model: {self.model_var.get()} | Language: {self.language_var.get()}", styles["Normal"]),
                        Spacer(1, 12),
                        Paragraph("Transcription:", styles["Heading2"]),
                        Spacer(1, 6)
                    ]
                    for line in self.full_transcription.split("\n"):
                        story.append(Paragraph(line, styles["Normal"]))
                        story.append(Spacer(1, 6))
                    doc.build(story)
                elif file_type == "docx":
                    doc = docx.Document()
                    doc.add_heading(f"Transcription of {os.path.basename(self.file_path)}", 0)
                    doc.add_paragraph(f"Model: {self.model_var.get()} | Language: {self.language_var.get()}")
                    doc.add_heading("Transcription:", level=2)
                    for line in self.full_transcription.split("\n"):
                        doc.add_paragraph(line)
                    doc.save(output_file)

                # Show success message with custom dialog
                success_dialog = tk.Toplevel(self.root)
                self.set_window_icon(success_dialog)
                success_dialog.title("Export")
                self.center_window(success_dialog, 830, 150)
                success_dialog.transient(self.root)
                success_dialog.grab_set()
                ttkb.Label(success_dialog, text=f"Transcription exported to '{output_file}'").pack(pady=10)
                ttkb.Button(success_dialog, text="OK", bootstyle="primary", command=success_dialog.destroy).pack(pady=10)
                self.status_var.set("Export Completed")
            except Exception as e:
                # Show error message with custom dialog
                error_dialog = tk.Toplevel(self.root)
                self.set_window_icon(error_dialog)
                error_dialog.title("Export Failed")
                self.center_window(error_dialog, 400, 150)
                error_dialog.transient(self.root)
                error_dialog.grab_set()
                ttkb.Label(error_dialog, text=f"Error during export: {e}").pack(pady=10)
                ttkb.Button(error_dialog, text="OK", bootstyle="primary", command=error_dialog.destroy).pack(pady=10)
                self.status_var.set("Export Failed")
            finally:
                self.is_processing = False
                self.spinner_label.pack_forget()
                self.update_playback_buttons()
                dialog.destroy()
        threading.Thread(target=do_export, daemon=True).start()

    def clear_output(self):
        self.stop_audio()
        self.output_text.delete(1.0, tk.END)
        self.analysis_output.delete(1.0, tk.END)
        self.file_path = None
        self.audio = None
        self.audio_duration = 0
        self.full_transcription = ""
        self.file_path_var.set("No file selected")
        self.duration_label.config(text="Duration: N/A")
        self.seek_scale.config(to=0)
        self.start_time_var.set(0)
        self.end_time_var.set(0)
        self.current_position = 0
        self.current_position_var.set("Current: 0s")
        self.selected_duration_var.set("Selected Duration: 0s")
        self.transcribe_button.config(state="disabled")
        self.play_button.config(state="disabled")
        self.stop_button.config(state="disabled")
        self.progress["value"] = 0
        self.status_var.set("Ready")
        self.update_playback_buttons()

if __name__ == "__main__":
    root = ttkb.Window()
    app = VoiceVoyager(root)
    root.mainloop()